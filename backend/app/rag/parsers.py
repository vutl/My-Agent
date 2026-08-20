from collections.abc import Callable
from dataclasses import dataclass, field, replace
import hashlib
import inspect
from io import BytesIO
from pathlib import Path
import re

from app.rag.figure_quality import (
    classify_visual_asset,
    extract_figure_label,
    group_visual_candidates,
    merge_visual_quality_metadata,
)


SUPPORTED_TEXT_TYPES = {"txt", "md", "pdf", "docx"}
VisionSummarizer = Callable[..., object | None]


@dataclass(frozen=True)
class ParsedPage:
    page_number: int
    text: str


@dataclass(frozen=True)
class ParsedTable:
    table_index: int
    page_number: int | None
    caption: str | None
    markdown: str | None
    row_count: int | None
    column_count: int | None
    extraction_method: str


@dataclass(frozen=True)
class ParsedFigure:
    figure_index: int
    page_number: int | None
    caption: str | None
    image_path: str | None
    visual_summary: str | None
    extraction_method: str
    bbox: dict[str, float] | None = None
    metadata: dict[str, object] = field(default_factory=dict)


@dataclass(frozen=True)
class ParsedDocument:
    text: str
    parser_name: str
    parser_version: str
    page_count: int | None = None
    pages: list[ParsedPage] | None = None
    tables: list[ParsedTable] | None = None
    figures: list[ParsedFigure] | None = None


def supported_file_type(path: Path) -> str | None:
    suffix = path.suffix.lower().lstrip(".")
    if suffix in SUPPORTED_TEXT_TYPES:
        return suffix
    return None


def parse_text_file(path: Path) -> str:
    return parse_document(path).text


def parse_document(
    path: Path,
    *,
    artifact_dir: Path | None = None,
    vision_summarizer: VisionSummarizer | None = None,
) -> ParsedDocument:
    file_type = supported_file_type(path)
    if file_type in {"txt", "md"}:
        text = path.read_text(encoding="utf-8", errors="replace")
        tables = _extract_markdown_tables(text) if file_type == "md" else []
        existing_table_captions = {table.caption for table in tables if table.caption}
        tables.extend(
            table
            for table in _caption_tables(text, page_number=1, start_index=len(tables))
            if table.caption not in existing_table_captions
        )
        figures = _caption_figures(text, page_number=1)
        return ParsedDocument(
            text=text,
            parser_name=f"{file_type}_plain_text",
            parser_version="1",
            page_count=1,
            pages=[ParsedPage(page_number=1, text=text)],
            tables=tables,
            figures=figures,
        )
    if file_type == "pdf":
        return _parse_pdf(path, artifact_dir=artifact_dir, vision_summarizer=vision_summarizer)
    if file_type == "docx":
        return _parse_docx(path)
    raise ValueError(f"Unsupported file type: {path.suffix}")


def _parse_pdf(
    path: Path,
    *,
    artifact_dir: Path | None = None,
    vision_summarizer: VisionSummarizer | None = None,
) -> ParsedDocument:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = []
    tables: list[ParsedTable] = []
    caption_figures: list[ParsedFigure] = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        pages.append(ParsedPage(page_number=index, text=text))
        tables.extend(_caption_tables(text, page_number=index, start_index=len(tables)))
        caption_figures.extend(_caption_figures(text, page_number=index, start_index=len(caption_figures)))

    if artifact_dir is not None:
        docling_parsed = _parse_pdf_with_docling(
            path=path,
            fallback_pages=pages,
            fallback_tables=tables,
            fallback_figures=caption_figures,
            artifact_dir=artifact_dir,
            vision_summarizer=vision_summarizer,
        )
        if docling_parsed is not None:
            return docling_parsed

    figures = _figures_with_artifacts(
        path=path,
        pages=pages,
        caption_figures=caption_figures,
        artifact_dir=artifact_dir,
        vision_summarizer=vision_summarizer,
    )
    figures = _with_page_visual_assets(
        path=path,
        artifact_dir=artifact_dir,
        pages=pages,
        tables=tables,
        figures=figures,
        caption_figures=caption_figures,
    )
    return ParsedDocument(
        text="\n\n".join(page.text for page in pages),
        parser_name="pypdf+pymupdf" if any(figure.image_path for figure in figures) else "pypdf",
        parser_version="3" if any(figure.image_path for figure in figures) else "1",
        page_count=len(reader.pages),
        pages=pages,
        tables=tables,
        figures=figures,
    )


def _parse_docx(path: Path) -> ParsedDocument:
    from docx import Document

    document = Document(str(path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    return ParsedDocument(
        text="\n".join(paragraphs),
        parser_name="python-docx",
        parser_version="1",
        page_count=None,
        pages=None,
        tables=[],
        figures=[],
    )


def _parse_pdf_with_docling(
    *,
    path: Path,
    fallback_pages: list[ParsedPage],
    fallback_tables: list[ParsedTable],
    fallback_figures: list[ParsedFigure],
    artifact_dir: Path,
    vision_summarizer: VisionSummarizer | None,
) -> ParsedDocument | None:
    try:
        from docling.datamodel.base_models import InputFormat  # type: ignore
        from docling.datamodel.pipeline_options import PdfPipelineOptions  # type: ignore
        from docling.document_converter import DocumentConverter, PdfFormatOption  # type: ignore
        from docling_core.types.doc import PictureItem, TableItem  # type: ignore
    except ImportError:
        return None

    try:
        pipeline_options = PdfPipelineOptions()
        # Image-only reports/slides otherwise produce zero text chunks and make
        # every downstream embedding/retrieval improvement irrelevant. Enable
        # OCR from source-text density, not from a filename or benchmark rule.
        ocr_enabled = _needs_pdf_ocr(fallback_pages)
        # Docling's OCR can improve item classification but, for some
        # image-only slide decks, those tokens are not exported into page text.
        # Keep structural parsing fast here and run the explicit page-text OCR
        # fallback below where provenance and output are under our control.
        pipeline_options.do_ocr = False
        pipeline_options.do_table_structure = True
        pipeline_options.images_scale = 2.0
        pipeline_options.generate_page_images = True
        pipeline_options.generate_picture_images = True
        converter = DocumentConverter(
            format_options={
                InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options)
            }
        )
        result = converter.convert(path)
        document = result.document
    except Exception:
        return None

    docling_text = _docling_markdown(document)
    text = docling_text or "\n\n".join(page.text for page in fallback_pages)
    pages = _docling_pages(document, fallback_pages)
    if ocr_enabled:
        pages = _ocr_sparse_pdf_pages(path, pages)
        text = "\n\n".join(page.text for page in pages)

    table_caption_offsets: dict[int, int] = {}
    table_captions_by_page: dict[int, list[str]] = {}
    for table in fallback_tables:
        if table.page_number is not None and table.caption:
            table_captions_by_page.setdefault(table.page_number, []).append(table.caption)

    tables: list[ParsedTable] = []
    for table in getattr(document, "tables", []) or []:
        page_number = _docling_page_number(table)
        markdown, row_count, column_count = _docling_table_markdown(table, document)
        caption = _docling_caption(table, document=document) or (
            _next_caption(table_captions_by_page, table_caption_offsets, page_number)
            if page_number is not None
            else None
        )
        tables.append(
            ParsedTable(
                table_index=len(tables),
                page_number=page_number,
                caption=caption,
                markdown=markdown,
                row_count=row_count,
                column_count=column_count,
                extraction_method="docling_table",
            )
        )
    if not tables:
        tables = fallback_tables

    figure_captions_by_page: dict[int, list[str]] = {}
    for figure in fallback_figures:
        if figure.page_number is not None and figure.caption:
            figure_captions_by_page.setdefault(figure.page_number, []).append(figure.caption)

    # PictureItems are deliberately collected before any regex/page-sequence
    # caption is assigned. A page caption belongs to a logical visual group,
    # not to whichever raw crop Docling happens to yield first.
    figures = _docling_logical_figures(
        path=path,
        document=document,
        picture_item_type=PictureItem,
        pages=pages,
        captions_by_page=figure_captions_by_page,
        artifact_dir=artifact_dir,
        vision_summarizer=vision_summarizer,
    )

    if not figures:
        figures = _figures_with_artifacts(
            path=path,
            pages=pages,
            caption_figures=fallback_figures,
            artifact_dir=artifact_dir,
            vision_summarizer=vision_summarizer,
        )
    figures = _with_page_visual_assets(
        path=path,
        artifact_dir=artifact_dir,
        pages=pages,
        tables=tables,
        figures=figures,
        caption_figures=fallback_figures,
    )

    return ParsedDocument(
        text=text,
        parser_name="docling",
        parser_version="4-page-ocr" if ocr_enabled else "4",
        page_count=len(pages) or None,
        pages=pages,
        tables=tables,
        figures=figures,
    )


def _needs_pdf_ocr(pages: list[ParsedPage]) -> bool:
    """Detect PDFs whose searchable text layer is too sparse to retrieve.

    The thresholds operate at document level so an ordinary paper with a few
    visual-only pages does not pay the cost of full OCR, while scanned reports
    and image-based slide decks do not silently index as empty documents.
    """

    if not pages:
        return False
    searchable_chars = [len(re.sub(r"\s+", "", page.text or "")) for page in pages]
    usable_pages = sum(count >= 40 for count in searchable_chars)
    usable_ratio = usable_pages / len(searchable_chars)
    mean_chars = sum(searchable_chars) / len(searchable_chars)
    return usable_ratio < 0.35 and mean_chars < 120


def _ocr_sparse_pdf_pages(path: Path, pages: list[ParsedPage]) -> list[ParsedPage]:
    """Populate sparse pages with full-page OCR while preserving page numbers.

    This is a local parser fallback. It never calls a language model and it
    leaves already-searchable pages untouched.
    """

    try:
        import fitz  # type: ignore
        from rapidocr import RapidOCR  # type: ignore
        from rapidocr.utils.typings import EngineType  # type: ignore
    except ImportError:
        return pages

    params = {
        f"{stage}.engine_type": EngineType.TORCH
        for stage in ("Det", "Cls", "Rec")
    }
    try:
        engine = RapidOCR(params=params)
    except Exception:
        return pages

    existing = {page.page_number: page for page in pages}
    updated: list[ParsedPage] = []
    try:
        with fitz.open(str(path)) as document:
            for page_number, pdf_page in enumerate(document, start=1):
                current = existing.get(page_number, ParsedPage(page_number, ""))
                if len(re.sub(r"\s+", "", current.text or "")) >= 40:
                    updated.append(current)
                    continue
                pixmap = pdf_page.get_pixmap(
                    matrix=fitz.Matrix(2.0, 2.0), alpha=False
                )
                result = engine(pixmap.tobytes("png"))
                texts = getattr(result, "txts", None) or ()
                ocr_text = "\n".join(
                    str(value).strip() for value in texts if str(value).strip()
                )
                updated.append(
                    ParsedPage(
                        page_number=page_number,
                        text=ocr_text or current.text,
                    )
                )
    except Exception:
        return pages
    return updated or pages


def _docling_markdown(document) -> str:
    for method_name in ("export_to_markdown", "export_to_text"):
        method = getattr(document, method_name, None)
        if method is None:
            continue
        try:
            text = method()
        except Exception:
            continue
        if isinstance(text, str) and text.strip():
            return text
    return ""


def _docling_pages(document, fallback_pages: list[ParsedPage]) -> list[ParsedPage]:
    pages: list[ParsedPage] = []
    doc_pages = getattr(document, "pages", None)
    if isinstance(doc_pages, dict):
        for page_no in sorted(doc_pages):
            page_text = ""
            export_method = getattr(doc_pages[page_no], "export_to_text", None)
            if export_method is not None:
                try:
                    page_text = export_method()
                except Exception:
                    page_text = ""
            if not page_text and page_no <= len(fallback_pages):
                page_text = fallback_pages[page_no - 1].text
            pages.append(ParsedPage(page_number=int(page_no), text=page_text))
    return pages or fallback_pages


def _docling_logical_figures(
    *,
    path: Path,
    document,
    picture_item_type,
    pages: list[ParsedPage],
    captions_by_page: dict[int, list[str]],
    artifact_dir: Path,
    vision_summarizer: VisionSummarizer | None,
) -> list[ParsedFigure]:
    """Turn Docling PictureItems into auditable logical figures.

    Docling may emit one PictureItem per sub-panel, or emit publisher marks as
    pictures. We retain those raw candidates as child diagnostics, but caption,
    classify, summarize and index only at logical-group granularity.
    """

    figures_dir = artifact_dir / "figures"
    children_dir = figures_dir / "children"
    figures_dir.mkdir(parents=True, exist_ok=True)
    children_dir.mkdir(parents=True, exist_ok=True)
    page_geometry = _pdf_page_geometry(path)

    candidates: list[dict[str, object]] = []
    seen_placements: set[tuple[object, ...]] = set()
    for source_index, (item, _level) in enumerate(document.iterate_items()):
        if not isinstance(item, picture_item_type):
            continue
        image = _docling_item_image(item, document)
        if image is None:
            continue

        image_bytes = _pil_png_bytes(image)
        digest = hashlib.sha256(image_bytes).hexdigest()
        page_number = _docling_page_number(item)
        bbox = _docling_bbox(item)
        placement_key = (
            page_number,
            digest,
            tuple(round(float((bbox or {}).get(key, 0.0)), 3) for key in ("x0", "y0", "x1", "y1")),
        )
        if placement_key in seen_placements:
            continue
        seen_placements.add(placement_key)

        direct_caption = _docling_caption(item, document=document)
        width = int(getattr(image, "width", 0) or 0)
        height = int(getattr(image, "height", 0) or 0)
        page_width, page_height = page_geometry.get(page_number, (None, None))
        child_path = children_dir / (
            f"page_{page_number or 0:03d}_item_{source_index + 1:03d}_{digest[:12]}.png"
        )
        child_path.write_bytes(image_bytes)
        candidates.append(
            {
                "source_index": source_index,
                "page_number": page_number,
                "caption": direct_caption,
                "bbox": bbox,
                "bbox_origin": _docling_bbox_origin(item),
                "image_hash": digest,
                "image_bytes": image_bytes,
                "image_path": str(child_path),
                "width": width,
                "height": height,
                "page_width": page_width,
                "page_height": page_height,
                "metadata": {
                    "image_hash": digest,
                    "caption_source": "docling_direct" if direct_caption else "none",
                },
            }
        )

    if not candidates:
        return []

    repeat_counts: dict[str, int] = {}
    digest_pages: dict[str, set[int | None]] = {}
    for candidate in candidates:
        digest_pages.setdefault(str(candidate["image_hash"]), set()).add(
            candidate.get("page_number") if isinstance(candidate.get("page_number"), int) else None
        )
    repeat_counts = {digest: len(page_numbers) for digest, page_numbers in digest_pages.items()}

    # Reject obvious repeated branding/header/footer assets before connected
    # component grouping, otherwise a logo close to a real panel can poison the
    # entire logical figure.
    eligible: list[dict[str, object]] = []
    rejected: list[dict[str, object]] = []
    for candidate in candidates:
        raw_metadata = _raw_candidate_metadata(candidate)
        decision = classify_visual_asset(
            caption=_optional_text(candidate.get("caption")),
            extraction_method="docling_picture",
            bbox=_optional_bbox(candidate.get("bbox")),
            metadata=raw_metadata,
            repeated_asset_count=repeat_counts.get(str(candidate["image_hash"]), 1),
        )
        candidate["raw_quality"] = decision.metadata_patch()
        if decision.status == "rejected":
            rejected.append(candidate)
        else:
            eligible.append(candidate)

    grouped_candidates: list[tuple[object, list[dict[str, object]]]] = []
    if eligible:
        for group in group_visual_candidates(eligible, document_key=str(path)):
            grouped_candidates.append((group, [eligible[index] for index in group.member_indices]))
    for candidate in rejected:
        singleton = group_visual_candidates([candidate], document_key=str(path))[0]
        grouped_candidates.append((singleton, [candidate]))

    grouped_candidates.sort(key=_logical_group_sort_key)
    captions_remaining = _remaining_fallback_captions(
        grouped_candidates=grouped_candidates,
        captions_by_page=captions_by_page,
    )

    figures: list[ParsedFigure] = []
    for group, members in grouped_candidates:
        page_number = getattr(group, "page_number", None)
        is_raw_rejected = all(
            (member.get("raw_quality") or {}).get("quality_status") == "rejected"
            for member in members
            if isinstance(member.get("raw_quality"), dict)
        )
        direct_caption = _best_direct_group_caption(members)
        if direct_caption:
            caption = direct_caption
            caption_source = "docling_direct"
        elif not is_raw_rejected:
            queue = captions_remaining.get(page_number, [])
            caption = queue.pop(0) if queue else None
            caption_source = "fallback_sequence" if caption else "none"
        else:
            caption = None
            caption_source = "none"

        logical_group_id = str(getattr(group, "logical_group_id"))
        group_bbox = _optional_bbox(getattr(group, "bbox", None))
        extraction_method = "docling_logical_composite" if len(members) > 1 else "docling_picture"
        image_path, image_width, image_height, composite_error = _logical_group_image(
            path=path,
            figures_dir=figures_dir,
            group_id=logical_group_id,
            page_number=page_number,
            bbox=group_bbox,
            members=members,
        )
        page_width = _first_number(member.get("page_width") for member in members)
        page_height = _first_number(member.get("page_height") for member in members)
        bbox_origin = _first_text(member.get("bbox_origin") for member in members) or "TOPLEFT"
        page_text = pages[page_number - 1].text if page_number and page_number <= len(pages) else ""
        child_diagnostics = [_child_diagnostic(member) for member in members]
        label = extract_figure_label(caption)
        metadata: dict[str, object] = {
            "source": "docling",
            "asset_type": "figure",
            "caption_source": caption_source,
            "logical_group_id": logical_group_id,
            "child_count": len(members),
            "children": child_diagnostics,
            "image_hash": _hash_path(image_path),
            "width": image_width,
            "height": image_height,
            "extension": "png",
            "bbox_origin": bbox_origin,
            "page_width": page_width,
            "page_height": page_height,
            "page_text_excerpt": page_text[:1800] if page_text else None,
        }
        if label:
            metadata["figure_label"] = label.label
            metadata["figure_number"] = label.number
        if composite_error:
            metadata["composite_render_error"] = composite_error

        repeated_count = max(
            repeat_counts.get(str(member.get("image_hash")), 1) for member in members
        )
        decision = classify_visual_asset(
            caption=caption,
            extraction_method=extraction_method,
            bbox=group_bbox,
            metadata=metadata,
            repeated_asset_count=repeated_count,
        )
        metadata.update(decision.metadata_patch())
        if composite_error and len(members) > 1:
            metadata.update(
                {
                    "quality_status": "needs_review",
                    "asset_kind": "panel",
                    "quality_confidence": 0.95,
                    "quality_reasons": [*list(metadata.get("quality_reasons") or []), "composite_render_failed"],
                    "is_complete": False,
                }
            )

        summary = None
        if vision_summarizer is not None and metadata.get("quality_status") != "rejected":
            page_image_path = _render_page_context_image(
                path=path,
                artifact_dir=artifact_dir,
                page_number=page_number,
            )
            try:
                summary, vision_patch = _summarize_visual(
                    vision_summarizer,
                    image_path=image_path,
                    caption=caption,
                    page_text=page_text,
                    page_number=page_number,
                    page_image_path=page_image_path,
                    extraction_method=extraction_method,
                )
                _merge_vision_quality(metadata, vision_patch)
            except Exception as exc:
                metadata["vision_error"] = str(exc)

        figures.append(
            ParsedFigure(
                figure_index=len(figures),
                page_number=page_number,
                caption=caption
                or (
                    f"Visual asset extracted from page {page_number}"
                    if page_number
                    else "Visual asset extracted by Docling"
                ),
                image_path=str(image_path),
                visual_summary=summary,
                extraction_method=extraction_method,
                bbox=group_bbox,
                metadata=metadata,
            )
        )

    return figures


def _raw_candidate_metadata(candidate: dict[str, object]) -> dict[str, object]:
    return {
        "width": candidate.get("width"),
        "height": candidate.get("height"),
        "page_width": candidate.get("page_width"),
        "page_height": candidate.get("page_height"),
        "caption_source": (candidate.get("metadata") or {}).get("caption_source")
        if isinstance(candidate.get("metadata"), dict)
        else "none",
    }


def _remaining_fallback_captions(
    *,
    grouped_candidates: list[tuple[object, list[dict[str, object]]]],
    captions_by_page: dict[int, list[str]],
) -> dict[int, list[str]]:
    remaining = {page: list(captions) for page, captions in captions_by_page.items()}
    for group, members in grouped_candidates:
        page_number = getattr(group, "page_number", None)
        direct = _best_direct_group_caption(members)
        direct_label = extract_figure_label(direct)
        if page_number is None or direct_label is None:
            continue
        page_captions = remaining.get(page_number) or []
        for index, fallback in enumerate(page_captions):
            fallback_label = extract_figure_label(fallback)
            if fallback_label and fallback_label.number == direct_label.number:
                page_captions.pop(index)
                break
    return remaining


def _best_direct_group_caption(members: list[dict[str, object]]) -> str | None:
    captions = [
        str(member["caption"]).strip()
        for member in members
        if isinstance(member.get("caption"), str) and str(member["caption"]).strip()
    ]
    if not captions:
        return None
    labelled = [caption for caption in captions if extract_figure_label(caption)]
    return max(labelled or captions, key=len)


def _logical_group_sort_key(
    item: tuple[object, list[dict[str, object]]],
) -> tuple[int, float, float, str]:
    group, _members = item
    bbox = _optional_bbox(getattr(group, "bbox", None)) or {
        "x0": 0.0,
        "y0": 0.0,
        "x1": 0.0,
        "y1": 0.0,
    }
    return (
        int(getattr(group, "page_number", 0) or 0),
        round(bbox["x0"], 2),
        -round(bbox["y1"], 2),
        str(getattr(group, "logical_group_id", "")),
    )


def _logical_group_image(
    *,
    path: Path,
    figures_dir: Path,
    group_id: str,
    page_number: int | None,
    bbox: dict[str, float] | None,
    members: list[dict[str, object]],
) -> tuple[Path, int, int, str | None]:
    if len(members) == 1:
        member = members[0]
        return (
            Path(str(member["image_path"])),
            int(member.get("width") or 0),
            int(member.get("height") or 0),
            None,
        )

    suffix = group_id.rsplit(":", 1)[-1]
    composite_path = figures_dir / f"page_{page_number or 0:03d}_logical_{suffix}.png"
    origin = _first_text(member.get("bbox_origin") for member in members) or "TOPLEFT"
    try:
        image_bytes, width, height = _render_pdf_bbox_crop(
            path=path,
            page_number=page_number,
            bbox=bbox,
            bbox_origin=origin,
        )
        composite_path.write_bytes(image_bytes)
        return composite_path, width, height, None
    except Exception as exc:
        # Preserve a debuggable record even if rendering the page crop fails;
        # the quality gate keeps this isolated panel out of retrieval.
        largest = max(
            members,
            key=lambda member: int(member.get("width") or 0) * int(member.get("height") or 0),
        )
        return (
            Path(str(largest["image_path"])),
            int(largest.get("width") or 0),
            int(largest.get("height") or 0),
            str(exc),
        )


def _render_pdf_bbox_crop(
    *,
    path: Path,
    page_number: int | None,
    bbox: dict[str, float] | None,
    bbox_origin: str,
) -> tuple[bytes, int, int]:
    if page_number is None or bbox is None:
        raise ValueError("logical group has no page/bbox")
    import fitz  # type: ignore

    with fitz.open(str(path)) as pdf:
        if page_number < 1 or page_number > len(pdf):
            raise ValueError(f"invalid page number: {page_number}")
        page = pdf[page_number - 1]
        page_rect = page.rect
        x0, x1 = sorted((bbox["x0"], bbox["x1"]))
        raw_y0, raw_y1 = sorted((bbox["y0"], bbox["y1"]))
        if bbox_origin.upper().endswith("BOTTOMLEFT"):
            y0 = page_rect.height - raw_y1
            y1 = page_rect.height - raw_y0
        else:
            y0, y1 = raw_y0, raw_y1
        padding = 6.0
        clip = fitz.Rect(x0 - padding, y0 - padding, x1 + padding, y1 + padding) & page_rect
        if clip.width < 8 or clip.height < 8:
            raise ValueError("logical group crop is empty")
        pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), clip=clip, alpha=False)
        return pixmap.tobytes("png"), pixmap.width, pixmap.height


def _render_page_context_image(
    *,
    path: Path,
    artifact_dir: Path,
    page_number: int | None,
) -> Path | None:
    if page_number is None:
        return None
    try:
        import fitz  # type: ignore

        pages_dir = artifact_dir / "pages"
        pages_dir.mkdir(parents=True, exist_ok=True)
        image_path = pages_dir / f"page_{page_number:03d}.png"
        if image_path.is_file():
            return image_path
        with fitz.open(str(path)) as pdf:
            if page_number < 1 or page_number > len(pdf):
                return None
            pixmap = pdf[page_number - 1].get_pixmap(
                matrix=fitz.Matrix(1.5, 1.5),
                alpha=False,
            )
            image_path.write_bytes(pixmap.tobytes("png"))
        return image_path
    except Exception:
        return None


def _summarize_visual(
    vision_summarizer: VisionSummarizer,
    *,
    image_path: Path,
    caption: str | None,
    page_text: str,
    page_number: int | None,
    page_image_path: Path | None,
    extraction_method: str,
) -> tuple[str | None, dict[str, object]]:
    kwargs = {
        "page_number": page_number,
        "page_image_path": page_image_path,
        "extraction_method": extraction_method,
    }
    try:
        signature = inspect.signature(vision_summarizer)
        accepts_context = any(
            parameter.kind == inspect.Parameter.VAR_KEYWORD
            for parameter in signature.parameters.values()
        ) or all(key in signature.parameters for key in kwargs)
    except (TypeError, ValueError):
        accepts_context = False

    result = (
        vision_summarizer(image_path, caption, page_text, **kwargs)
        if accepts_context
        else vision_summarizer(image_path, caption, page_text)
    )
    if result is None:
        return None, {}
    raw_text = getattr(result, "raw_text", None)
    metadata_method = getattr(result, "to_metadata", None)
    if isinstance(raw_text, str):
        patch = metadata_method() if callable(metadata_method) else {}
        return raw_text, patch if isinstance(patch, dict) else {}
    return str(result), {}


def _merge_vision_quality(metadata: dict[str, object], patch: dict[str, object]) -> None:
    merge_visual_quality_metadata(metadata, patch)


def _child_diagnostic(candidate: dict[str, object]) -> dict[str, object]:
    return {
        "source_index": candidate.get("source_index"),
        "image_hash": candidate.get("image_hash"),
        "image_path": candidate.get("image_path"),
        "bbox": candidate.get("bbox"),
        "bbox_origin": candidate.get("bbox_origin"),
        "width": candidate.get("width"),
        "height": candidate.get("height"),
        "direct_caption": candidate.get("caption"),
        **(
            candidate.get("raw_quality")
            if isinstance(candidate.get("raw_quality"), dict)
            else {}
        ),
    }


def _pdf_page_geometry(path: Path) -> dict[int, tuple[float | None, float | None]]:
    try:
        import fitz  # type: ignore

        with fitz.open(str(path)) as pdf:
            return {
                page_number: (float(page.rect.width), float(page.rect.height))
                for page_number, page in enumerate(pdf, start=1)
            }
    except Exception:
        return {}


def _optional_bbox(value: object) -> dict[str, float] | None:
    if not isinstance(value, dict):
        return None
    try:
        return {key: float(value[key]) for key in ("x0", "y0", "x1", "y1")}
    except (KeyError, TypeError, ValueError):
        return None


def _optional_text(value: object) -> str | None:
    return value.strip() if isinstance(value, str) and value.strip() else None


def _first_text(values) -> str | None:
    return next((str(value) for value in values if isinstance(value, str) and value), None)


def _first_number(values) -> float | None:
    for value in values:
        if isinstance(value, int | float) and float(value) > 0:
            return float(value)
    return None


def _hash_path(path: Path) -> str | None:
    try:
        return hashlib.sha256(path.read_bytes()).hexdigest()
    except OSError:
        return None


def _docling_table_markdown(table, document) -> tuple[str | None, int | None, int | None]:
    try:
        dataframe = table.export_to_dataframe(doc=document)
        row_count = len(dataframe.index)
        column_count = len(dataframe.columns)
        try:
            markdown = dataframe.to_markdown(index=False)
        except Exception:
            markdown = dataframe.to_csv(index=False)
        return markdown, row_count, column_count
    except Exception:
        pass

    try:
        html = table.export_to_html(doc=document)
    except Exception:
        html = None
    return html, None, None


def _docling_item_image(item, document):
    get_image = getattr(item, "get_image", None)
    if get_image is None:
        return None
    try:
        return get_image(document)
    except Exception:
        return None


def _pil_png_bytes(image) -> bytes:
    buffer = BytesIO()
    image.save(buffer, format="PNG")
    return buffer.getvalue()


def _skip_docling_picture(*, width: int, height: int, caption: str | None) -> bool:
    has_figure_caption = bool(caption and _caption_re("figure").match(caption.strip()))
    if has_figure_caption:
        return False
    if width < 160 or height < 120:
        return True
    if width * height < 25_000:
        return True
    return False


def _docling_caption(item, *, document=None) -> str | None:
    # Current Docling exposes caption_text(doc); older versions exposed string
    # attributes. Calling the method without the document silently lost every
    # direct caption and forced unsafe page-sequence assignment.
    for attr in ("caption_text", "caption", "text", "name"):
        value = getattr(item, attr, None)
        if callable(value):
            try:
                value = value(document) if document is not None else value()
            except Exception:
                value = None
        if isinstance(value, str) and value.strip():
            return " ".join(value.split())[:500]
    return None


def _docling_page_number(item) -> int | None:
    prov = getattr(item, "prov", None)
    if isinstance(prov, list) and prov:
        page_no = getattr(prov[0], "page_no", None)
        if isinstance(page_no, int):
            return page_no
    page_no = getattr(item, "page_no", None)
    return page_no if isinstance(page_no, int) else None


def _docling_bbox(item) -> dict[str, float] | None:
    prov = getattr(item, "prov", None)
    if not isinstance(prov, list) or not prov:
        return None
    bbox = getattr(prov[0], "bbox", None)
    if bbox is None:
        return None
    values = {}
    for attr in ("l", "t", "r", "b", "x0", "y0", "x1", "y1"):
        value = getattr(bbox, attr, None)
        if isinstance(value, int | float):
            values[attr] = float(value)
    if {"l", "t", "r", "b"}.issubset(values):
        return {"x0": values["l"], "y0": values["t"], "x1": values["r"], "y1": values["b"]}
    if {"x0", "y0", "x1", "y1"}.issubset(values):
        return {"x0": values["x0"], "y0": values["y0"], "x1": values["x1"], "y1": values["y1"]}
    return None


def _docling_bbox_origin(item) -> str:
    prov = getattr(item, "prov", None)
    if not isinstance(prov, list) or not prov:
        return "TOPLEFT"
    bbox = getattr(prov[0], "bbox", None)
    origin = getattr(bbox, "coord_origin", None)
    value = getattr(origin, "value", origin)
    return str(value or "TOPLEFT")


def _with_page_visual_assets(
    *,
    path: Path,
    artifact_dir: Path | None,
    pages: list[ParsedPage],
    tables: list[ParsedTable],
    figures: list[ParsedFigure],
    caption_figures: list[ParsedFigure],
) -> list[ParsedFigure]:
    if artifact_dir is None:
        return figures

    page_assets = _render_pdf_page_assets(
        path=path,
        artifact_dir=artifact_dir,
        pages=pages,
        tables=tables,
        figures=figures,
        caption_figures=caption_figures,
    )
    return [
        replace(figure, figure_index=index)
        for index, figure in enumerate([*figures, *page_assets])
    ]


def _render_pdf_page_assets(
    *,
    path: Path,
    artifact_dir: Path,
    pages: list[ParsedPage],
    tables: list[ParsedTable],
    figures: list[ParsedFigure],
    caption_figures: list[ParsedFigure],
) -> list[ParsedFigure]:
    try:
        import fitz  # type: ignore
    except ImportError:
        return []

    pages_dir = artifact_dir / "pages"
    pages_dir.mkdir(parents=True, exist_ok=True)
    visual_captions_by_page: dict[int, list[str]] = {}
    for table in tables:
        if table.page_number is not None and table.caption:
            visual_captions_by_page.setdefault(table.page_number, []).append(table.caption)
    for figure in caption_figures:
        if figure.page_number is not None and figure.caption:
            visual_captions_by_page.setdefault(figure.page_number, []).append(figure.caption)

    accepted_figure_pages = {
        figure.page_number
        for figure in figures
        if figure.page_number is not None
        and figure.image_path
        and (figure.metadata or {}).get("asset_type", "figure") != "page"
    }

    page_assets: list[ParsedFigure] = []
    with fitz.open(str(path)) as document:
        for page_number, page in enumerate(document, start=1):
            pixmap = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
            image_path = pages_dir / f"page_{page_number:03d}.png"
            image_path.write_bytes(pixmap.tobytes("png"))
            captions = visual_captions_by_page.get(page_number, [])
            drawing_count = _page_drawing_count(page)
            should_index_page = (
                page_number not in accepted_figure_pages
                and bool(captions)
            ) or drawing_count >= 40
            if not should_index_page:
                continue

            page_text = pages[page_number - 1].text if page_number <= len(pages) else ""
            page_assets.append(
                ParsedFigure(
                    figure_index=len(figures) + len(page_assets),
                    page_number=page_number,
                    caption=f"Page {page_number} visual fallback",
                    image_path=str(image_path),
                    visual_summary=None,
                    extraction_method="page_screenshot",
                    bbox=_rect_dict(page.rect),
                    metadata={
                        "source": "pymupdf_page_render",
                        "asset_type": "page",
                        "status": "fallback",
                        "quality_status": "rejected",
                        "asset_kind": "page",
                        "quality_confidence": 0.99,
                        "quality_reasons": ["page_render_fallback"],
                        "is_content": False,
                        "is_complete": True,
                        "width": pixmap.width,
                        "height": pixmap.height,
                        "extension": "png",
                        "detected_captions": captions[:8],
                        "drawing_count": drawing_count,
                        "page_text_excerpt": page_text[:1200],
                    },
                )
            )
    return page_assets


def _page_drawing_count(page) -> int:
    try:
        return len(page.get_drawings())
    except Exception:
        return 0


def _extract_markdown_tables(text: str) -> list[ParsedTable]:
    lines = text.splitlines()
    tables: list[ParsedTable] = []
    index = 0
    while index < len(lines):
        if not _looks_like_markdown_table_start(lines, index):
            index += 1
            continue

        start = index
        index += 2
        while index < len(lines) and "|" in lines[index]:
            index += 1

        table_lines = lines[start:index]
        header = _split_markdown_row(table_lines[0])
        rows = [
            _split_markdown_row(line)
            for line in table_lines[2:]
            if _split_markdown_row(line)
        ]
        tables.append(
            ParsedTable(
                table_index=len(tables),
                page_number=1,
                caption=_previous_caption(lines, start, "table"),
                markdown="\n".join(table_lines),
                row_count=len(rows),
                column_count=len(header) if header else None,
                extraction_method="markdown_pipe_table",
            )
        )
    return tables


def _looks_like_markdown_table_start(lines: list[str], index: int) -> bool:
    if index + 1 >= len(lines):
        return False
    return "|" in lines[index] and bool(re.search(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", lines[index + 1]))


def _split_markdown_row(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    return [cell.strip() for cell in stripped.split("|") if cell.strip()]


def _previous_caption(lines: list[str], start_index: int, kind: str) -> str | None:
    for line in reversed(lines[max(0, start_index - 3):start_index]):
        stripped = line.strip()
        if _caption_re(kind).match(stripped):
            return stripped
    return None


def _caption_tables(text: str, *, page_number: int | None, start_index: int = 0) -> list[ParsedTable]:
    return [
        ParsedTable(
            table_index=start_index + index,
            page_number=page_number,
            caption=caption,
            markdown=None,
            row_count=None,
            column_count=None,
            extraction_method="caption_regex",
        )
        for index, caption in enumerate(_caption_lines(text, "table"))
    ]


def _caption_figures(text: str, *, page_number: int | None, start_index: int = 0) -> list[ParsedFigure]:
    return [
        ParsedFigure(
            figure_index=start_index + index,
            page_number=page_number,
            caption=caption,
            image_path=None,
            visual_summary=None,
            extraction_method="caption_regex",
        )
        for index, caption in enumerate(_caption_lines(text, "figure"))
    ]


def _caption_lines(text: str, kind: str) -> list[str]:
    pattern = _caption_re(kind)
    other_pattern = _caption_re("table" if kind == "figure" else "figure")
    captions: list[str] = []
    current: list[str] = []

    def flush() -> None:
        if not current:
            return
        captions.append(_join_caption_parts(current)[:500])
        current.clear()

    for line in text.splitlines():
        stripped = " ".join(line.split())
        if not stripped:
            flush()
            continue
        if pattern.match(stripped):
            flush()
            current.append(stripped)
            continue
        if other_pattern.match(stripped):
            flush()
            continue
        if current:
            # A Markdown pipe table starts immediately after its caption. Its
            # rows are table data, not caption continuation text.
            if kind == "table" and "|" in stripped:
                flush()
                continue
            current.append(stripped)

    flush()
    return captions


def _join_caption_parts(parts: list[str]) -> str:
    if not parts:
        return ""
    merged = parts[0]
    for part in parts[1:]:
        if merged.rstrip().endswith("-") and part and part[0].islower():
            merged = merged.rstrip()[:-1] + part
        else:
            merged = f"{merged.rstrip()} {part.lstrip()}"
    return " ".join(merged.split())


def _caption_re(kind: str) -> re.Pattern[str]:
    if kind == "table":
        return re.compile(r"^(table|bảng)\s*\d+\s*[\.:–-]", re.IGNORECASE)
    return re.compile(r"^(fig\.?|figure|hình)\s*\d+\s*[\.:–-]", re.IGNORECASE)


def _figures_with_artifacts(
    *,
    path: Path,
    pages: list[ParsedPage],
    caption_figures: list[ParsedFigure],
    artifact_dir: Path | None,
    vision_summarizer: VisionSummarizer | None,
) -> list[ParsedFigure]:
    if artifact_dir is None:
        return caption_figures

    caption_crops = _extract_pdf_caption_crops(
        path=path,
        artifact_dir=artifact_dir,
        pages=pages,
        caption_figures=caption_figures,
        vision_summarizer=vision_summarizer,
    )
    used_captions = {
        (figure.page_number, figure.caption)
        for figure in caption_crops
        if figure.caption
    }
    remaining_caption_figures = [
        figure
        for figure in caption_figures
        if (figure.page_number, figure.caption) not in used_captions
    ]

    captions_by_page: dict[int, list[str]] = {}
    for figure in remaining_caption_figures:
        if figure.page_number is not None and figure.caption:
            captions_by_page.setdefault(figure.page_number, []).append(figure.caption)

    extracted_images = _extract_pdf_image_figures(
        path=path,
        artifact_dir=artifact_dir,
        pages=pages,
        captions_by_page=captions_by_page,
        vision_summarizer=vision_summarizer,
        skip_pages={figure.page_number for figure in caption_crops if figure.page_number is not None},
    )
    used_captions = {
        (figure.page_number, figure.caption)
        for figure in [*caption_crops, *extracted_images]
        if figure.caption
    }
    remaining_caption_figures = [
        figure
        for figure in caption_figures
        if (figure.page_number, figure.caption) not in used_captions
    ]
    return [
        replace(figure, figure_index=index)
        for index, figure in enumerate([*caption_crops, *extracted_images, *remaining_caption_figures])
    ]


def _extract_pdf_caption_crops(
    *,
    path: Path,
    artifact_dir: Path,
    pages: list[ParsedPage],
    caption_figures: list[ParsedFigure],
    vision_summarizer: VisionSummarizer | None,
) -> list[ParsedFigure]:
    try:
        import fitz  # type: ignore
    except ImportError:
        return []

    figures_dir = artifact_dir / "figures"
    figures_dir.mkdir(parents=True, exist_ok=True)
    figures: list[ParsedFigure] = []
    seen_hashes: set[str] = set()

    with fitz.open(str(path)) as document:
        for caption_figure in caption_figures:
            if caption_figure.page_number is None or not caption_figure.caption:
                continue
            page_index = caption_figure.page_number
            if page_index < 1 or page_index > len(document):
                continue
            page = document[page_index - 1]
            caption_rect = _find_caption_rect(page, caption_figure.caption)
            if caption_rect is None:
                continue
            crop_rect = _caption_figure_crop_rect(page, caption_rect)
            if crop_rect is None:
                continue

            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), clip=crop_rect, alpha=False)
            if pixmap.width < 96 or pixmap.height < 96:
                continue
            image_bytes = pixmap.tobytes("png")
            digest = hashlib.sha256(image_bytes).hexdigest()
            if digest in seen_hashes:
                continue
            seen_hashes.add(digest)

            image_path = figures_dir / f"page_{page_index:03d}_figure_{len(figures) + 1:03d}_{digest[:12]}.png"
            image_path.write_bytes(image_bytes)

            summary = None
            page_text = pages[page_index - 1].text if page_index <= len(pages) else ""
            metadata: dict[str, object] = {
                "source": "pymupdf_caption_crop",
                "asset_type": "figure",
                "image_hash": digest,
                "width": pixmap.width,
                "height": pixmap.height,
                "extension": "png",
                "caption_source": "caption_crop",
                "page_width": float(page.rect.width),
                "page_height": float(page.rect.height),
                "bbox_origin": "TOPLEFT",
                "caption_bbox": _rect_dict(caption_rect),
                "page_text_excerpt": page_text[:600] if page_text else None,
            }
            label = extract_figure_label(caption_figure.caption)
            if label:
                metadata["figure_label"] = label.label
                metadata["figure_number"] = label.number
            metadata.update(
                classify_visual_asset(
                    caption=caption_figure.caption,
                    extraction_method="pymupdf_caption_crop",
                    bbox=_rect_dict(crop_rect),
                    metadata=metadata,
                ).metadata_patch()
            )
            if vision_summarizer is not None:
                try:
                    summary, vision_patch = _summarize_visual(
                        vision_summarizer,
                        image_path=image_path,
                        caption=caption_figure.caption,
                        page_text=page_text,
                        page_number=page_index,
                        page_image_path=_render_page_context_image(
                            path=path,
                            artifact_dir=artifact_dir,
                            page_number=page_index,
                        ),
                        extraction_method="pymupdf_caption_crop",
                    )
                    _merge_vision_quality(metadata, vision_patch)
                except Exception as exc:
                    metadata["vision_error"] = str(exc)

            figures.append(
                ParsedFigure(
                    figure_index=len(figures),
                    page_number=page_index,
                    caption=caption_figure.caption,
                    image_path=str(image_path),
                    visual_summary=summary,
                    extraction_method="pymupdf_caption_crop",
                    bbox=_rect_dict(crop_rect),
                    metadata=metadata,
                )
            )

    return figures


def _extract_pdf_image_figures(
    *,
    path: Path,
    artifact_dir: Path,
    pages: list[ParsedPage],
    captions_by_page: dict[int, list[str]],
    vision_summarizer: VisionSummarizer | None,
    skip_pages: set[int] | None = None,
) -> list[ParsedFigure]:
    try:
        import fitz  # type: ignore
    except ImportError:
        return []

    figures_dir = artifact_dir / "figures"
    figures_dir.mkdir(parents=True, exist_ok=True)
    caption_offsets: dict[int, int] = {}
    seen_hashes: set[str] = set()
    figures: list[ParsedFigure] = []

    with fitz.open(str(path)) as document:
        for page_index, page in enumerate(document, start=1):
            if skip_pages and page_index in skip_pages:
                continue
            page_text = pages[page_index - 1].text if page_index <= len(pages) else ""
            for image_ref in page.get_images(full=True):
                xref = image_ref[0]
                extracted = document.extract_image(xref)
                image_bytes = extracted.get("image")
                if not image_bytes:
                    continue

                width = int(extracted.get("width") or 0)
                height = int(extracted.get("height") or 0)
                if width < 32 or height < 32:
                    continue

                digest = hashlib.sha256(image_bytes).hexdigest()
                if digest in seen_hashes:
                    continue
                seen_hashes.add(digest)

                extension = _normalized_image_extension(str(extracted.get("ext") or "png"))
                image_path = figures_dir / f"page_{page_index:03d}_figure_{len(figures) + 1:03d}_{digest[:12]}.{extension}"
                image_path.write_bytes(image_bytes)

                caption = _next_caption(captions_by_page, caption_offsets, page_index)
                summary = None
                metadata: dict[str, object] = {
                    "source": "pymupdf",
                    "asset_type": "figure",
                    "image_hash": digest,
                    "width": width,
                    "height": height,
                    "extension": extension,
                    "caption_source": "fallback_sequence" if caption else "none",
                    "page_width": float(page.rect.width),
                    "page_height": float(page.rect.height),
                    "bbox_origin": "TOPLEFT",
                    "page_text_excerpt": page_text[:600] if page_text else None,
                }
                label = extract_figure_label(caption)
                if label:
                    metadata["figure_label"] = label.label
                    metadata["figure_number"] = label.number
                image_bbox = _first_image_bbox(page, xref)
                metadata.update(
                    classify_visual_asset(
                        caption=caption,
                        extraction_method="pymupdf_image",
                        bbox=image_bbox,
                        metadata=metadata,
                    ).metadata_patch()
                )
                if vision_summarizer is not None:
                    try:
                        summary, vision_patch = _summarize_visual(
                            vision_summarizer,
                            image_path=image_path,
                            caption=caption,
                            page_text=page_text,
                            page_number=page_index,
                            page_image_path=_render_page_context_image(
                                path=path,
                                artifact_dir=artifact_dir,
                                page_number=page_index,
                            ),
                            extraction_method="pymupdf_image",
                        )
                        _merge_vision_quality(metadata, vision_patch)
                    except Exception as exc:
                        metadata["vision_error"] = str(exc)

                figures.append(
                    ParsedFigure(
                        figure_index=len(figures),
                        page_number=page_index,
                        caption=caption or f"Figure extracted from page {page_index}",
                        image_path=str(image_path),
                        visual_summary=summary,
                        extraction_method="pymupdf_image",
                        bbox=image_bbox,
                        metadata=metadata,
                    )
                )

    return figures


def _next_caption(
    captions_by_page: dict[int, list[str]],
    caption_offsets: dict[int, int],
    page_number: int,
) -> str | None:
    captions = captions_by_page.get(page_number) or []
    offset = caption_offsets.get(page_number, 0)
    caption_offsets[page_number] = offset + 1
    if offset >= len(captions):
        return None
    return captions[offset]


def _first_image_bbox(page, xref: int) -> dict[str, float] | None:
    try:
        rects = page.get_image_rects(xref)
    except Exception:
        return None
    if not rects:
        return None
    rect = rects[0]
    return _rect_dict(rect)


def _find_caption_rect(page, caption: str):
    variants = _caption_search_variants(caption)
    for variant in variants:
        try:
            rects = page.search_for(variant)
        except Exception:
            rects = []
        if rects:
            rect = rects[0]
            for extra in rects[1:4]:
                rect |= extra
            return rect
    return None


def _caption_search_variants(caption: str) -> list[str]:
    normalized = " ".join(caption.split())
    variants = [normalized]
    match = re.match(r"^((?:fig\.?|figure|hình)\s*\d+)", normalized, flags=re.IGNORECASE)
    if match:
        variants.append(match.group(1))
    if ":" in normalized:
        variants.append(normalized.split(":", 1)[0])
    if "." in normalized:
        variants.append(normalized.split(".", 1)[0])
    variants.append(normalized[:80])
    seen: set[str] = set()
    unique: list[str] = []
    for variant in variants:
        text = variant.strip()
        if len(text) >= 4 and text.lower() not in seen:
            seen.add(text.lower())
            unique.append(text)
    return unique


def _caption_figure_crop_rect(page, caption_rect):
    try:
        import fitz  # type: ignore
    except ImportError:
        return None

    page_rect = page.rect
    margin_x = 36
    top_margin = 48
    bottom = max(page_rect.y0 + top_margin + 80, caption_rect.y0 - 8)
    crop_height = min(330, max(140, caption_rect.y0 - page_rect.y0 - top_margin))
    top = max(page_rect.y0 + top_margin, bottom - crop_height)
    if bottom - top < 80:
        top = caption_rect.y1 + 8
        bottom = min(page_rect.y1 - 48, top + 300)
    if bottom - top < 80:
        return None
    return fitz.Rect(
        page_rect.x0 + margin_x,
        top,
        page_rect.x1 - margin_x,
        bottom,
    )


def _rect_dict(rect) -> dict[str, float]:
    return {
        "x0": float(rect.x0),
        "y0": float(rect.y0),
        "x1": float(rect.x1),
        "y1": float(rect.y1),
    }


def _normalized_image_extension(extension: str) -> str:
    normalized = extension.lower().lstrip(".")
    if normalized == "jpeg":
        return "jpg"
    if not re.fullmatch(r"[a-z0-9]+", normalized):
        return "png"
    return normalized or "png"
