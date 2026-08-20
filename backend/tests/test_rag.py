from pathlib import Path

from app.db.sqlite import connect, init_db
from app.rag.chunking import chunk_parsed_document, chunk_text, count_tokens
from app.rag.parsers import ParsedDocument, ParsedPage, parse_document, parse_text_file
from app.services.indexing_service import IndexingService
from app.services.rag_service import RagService


def test_chunk_text_splits_long_text() -> None:
    text = "alpha " * 500

    chunks = chunk_text(text, chunk_size=200, overlap=20)

    assert len(chunks) > 1
    assert all(chunk for chunk in chunks)


def test_chunk_parsed_document_preserves_page_numbers() -> None:
    parsed = ParsedDocument(
        text="first page text\n\nsecond page text",
        parser_name="test",
        parser_version="1",
        page_count=2,
        pages=[
            ParsedPage(page_number=7, text="first page text"),
            ParsedPage(page_number=8, text="second page text"),
        ],
    )

    chunks = chunk_parsed_document(parsed, chunk_size=200, overlap=20)

    assert [chunk.page_number for chunk in chunks] == [7, 8]
    assert [chunk.text for chunk in chunks] == ["first page text", "second page text"]


def test_chunking_preserves_paragraph_boundaries_and_section_title() -> None:
    text = "\n\n".join(
        [
            "2 Method",
            "The first paragraph explains the architecture. It should stay together.",
            "The second paragraph adds training details. It should not be cut in the middle.",
        ]
    )

    chunks = chunk_parsed_document(
        ParsedDocument(text=text, parser_name="test", parser_version="1"),
        chunk_size=120,
        overlap=0,
    )

    assert chunks
    assert chunks[0].section_title == "2 Method"
    assert any("The first paragraph explains the architecture." in chunk.text for chunk in chunks)


def test_token_chunking_creates_bounded_children_with_parent_context() -> None:
    text = "\n\n".join(
        [
            "2 Method",
            "The architecture combines acoustic and visual encoders. " * 35,
            "The fusion block exchanges information across modalities. " * 35,
        ]
    )

    chunks = chunk_parsed_document(
        ParsedDocument(text=text, parser_name="test", parser_version="1"),
        max_tokens=96,
        overlap_tokens=12,
        parent_max_tokens=320,
    )

    assert len(chunks) > 1
    assert all((chunk.token_count or 0) <= 96 for chunk in chunks)
    assert all(chunk.parent_text for chunk in chunks)
    assert all((chunk.parent_token_count or 0) <= 320 for chunk in chunks)
    assert any(chunk.parent_text != chunk.text for chunk in chunks)
    assert all(count_tokens(chunk.text) == chunk.token_count for chunk in chunks)


def test_markdown_table_and_figure_caption_are_indexed(tmp_path) -> None:
    db_path = tmp_path / "app.db"
    docs_dir = tmp_path / "docs"
    docs_dir.mkdir()
    (docs_dir / "report.md").write_text(
        "\n".join(
            [
                "# Report",
                "",
                "Table 1: Metrics",
                "| metric | value |",
                "| --- | --- |",
                "| accuracy | 0.92 |",
                "",
                "Figure 1: Agent architecture diagram",
            ]
        ),
        encoding="utf-8",
    )

    init_db(db_path)
    indexing = IndexingService(db_path)
    rag = RagService(db_path)
    result = indexing.index_folder(
        folder_path=str(docs_dir),
        recursive=False,
        file_types=["md"],
    )
    document_id = result.documents[0].id
    documents = indexing.list_documents()

    assert documents[0]["table_count"] == 1
    assert documents[0]["figure_count"] == 1

    tables = rag.list_document_tables(document_id)
    figures = rag.list_document_figures(document_id)

    assert tables[0]["caption"] == "Table 1: Metrics"
    assert tables[0]["row_count"] == 1
    assert tables[0]["column_count"] == 2
    assert figures[0]["caption"] == "Figure 1: Agent architecture diagram"


def test_mislabeled_table_caption_is_reconciled_from_unique_same_page_text(tmp_path) -> None:
    db_path = tmp_path / "app.db"
    source = tmp_path / "report.md"
    source.write_text(
        """# Report

Table 5: Existing results
| Model | F1 |
| --- | --- |
| Base | 0.50 |

Table 6: Target results
| Model | F1 |
| --- | --- |
| Ours | 0.80 |
""",
        encoding="utf-8",
    )
    init_db(db_path)
    indexed = IndexingService(db_path).index_file(source_path=str(source))
    with connect(db_path) as connection:
        connection.execute(
            """
            UPDATE document_tables
            SET caption = 'Figure 5: Nearby qualitative example'
            WHERE document_id = ? AND table_index = 1
            """,
            (indexed.id,),
        )

    tables = RagService(db_path).list_document_tables(indexed.id)

    assert [table["caption"] for table in tables] == [
        "Table 5: Existing results",
        "Table 6: Target results",
    ]
    repaired = tables[1]
    assert repaired["metadata"]["caption_reconciled"] is True
    assert repaired["metadata"]["caption_reconciled_number"] == 6
    assert (
        repaired["metadata"]["caption_reconciliation_source"]
        == "same_document_page_text"
    )


def test_mislabeled_table_caption_stays_filtered_when_page_match_is_ambiguous(tmp_path) -> None:
    db_path = tmp_path / "app.db"
    source = tmp_path / "report.md"
    source.write_text(
        """# Report

Table 5: Existing results
| Model | F1 |
| --- | --- |
| Base | 0.50 |

Table 6: First unmatched caption
| Model | F1 |
| --- | --- |
| Ours | 0.80 |

Table 7: Second unmatched caption
| Model | F1 |
| --- | --- |
| Ours | 0.81 |
""",
        encoding="utf-8",
    )
    init_db(db_path)
    indexed = IndexingService(db_path).index_file(source_path=str(source))
    with connect(db_path) as connection:
        connection.execute(
            """
            UPDATE document_tables
            SET caption = 'Figure 5: Nearby qualitative example'
            WHERE document_id = ? AND table_index = 1
            """,
            (indexed.id,),
        )
        connection.execute(
            "DELETE FROM document_tables WHERE document_id = ? AND table_index = 2",
            (indexed.id,),
        )

    tables = RagService(db_path).list_document_tables(indexed.id)

    assert [table["caption"] for table in tables] == ["Table 5: Existing results"]


def test_mislabeled_table_caption_does_not_borrow_other_page_text(tmp_path) -> None:
    db_path = tmp_path / "app.db"
    source = tmp_path / "report.md"
    source.write_text(
        """# Report

Table 6: Target results
| Model | F1 |
| --- | --- |
| Ours | 0.80 |
""",
        encoding="utf-8",
    )
    init_db(db_path)
    indexed = IndexingService(db_path).index_file(source_path=str(source))
    with connect(db_path) as connection:
        connection.execute(
            """
            UPDATE document_tables
            SET caption = 'Figure 5: Nearby qualitative example', page_number = 2
            WHERE document_id = ?
            """,
            (indexed.id,),
        )

    assert RagService(db_path).list_document_tables(indexed.id) == []


def test_index_folder_and_search_markdown(tmp_path) -> None:
    db_path = tmp_path / "app.db"
    docs_dir = tmp_path / "docs"
    docs_dir.mkdir()
    (docs_dir / "project.md").write_text(
        "# Project\n\nPhase three adds local document retrieval for planning.",
        encoding="utf-8",
    )

    init_db(db_path)
    indexing = IndexingService(db_path)
    rag = RagService(db_path)

    result = indexing.index_folder(
        folder_path=str(docs_dir),
        recursive=False,
        file_types=["md"],
    )
    chunks = rag.search("document retrieval", top_k=3)

    assert result.scanned_files == 1
    assert result.indexed_files == 1
    assert result.documents[0].filename == "project.md"
    assert chunks
    assert chunks[0].filename == "project.md"
    assert chunks[0].chunk_index == 0
    assert chunks[0].page_number == 1
    assert chunks[0].token_count is not None

    documents = indexing.list_documents()
    assert documents[0]["parser_name"] == "md_plain_text"
    assert documents[0]["index_status"] == "indexed"

    detail = rag.get_document(documents[0]["id"])
    assert detail is not None
    assert detail["filename"] == "project.md"

    stored_chunks = rag.list_document_chunks(documents[0]["id"])
    assert stored_chunks[0]["page_number"] == 1
    assert stored_chunks[0]["metadata"]["file_type"] == "md"
    assert stored_chunks[0]["metadata"]["page_number"] == 1
    assert stored_chunks[0]["metadata"]["context_prefix"].startswith("Document: project.md.")

    debug = rag.search_debug("document retrieval", top_k=3)
    assert debug["fts_query"]
    assert debug["retrieval_channels"] == ["sqlite_fts5"]


def test_version_reindex_failure_keeps_previous_canonical_document(tmp_path, monkeypatch) -> None:
    import json
    import pytest

    from app.db.sqlite import connect
    from app.services import indexing_service as indexing_module

    db_path = tmp_path / "app.db"
    artifact_root = tmp_path / "artifacts"
    source = tmp_path / "paper.md"
    source.write_text("# Paper\n\nStable evidence.", encoding="utf-8")
    init_db(db_path)
    indexing = IndexingService(db_path, artifact_root=artifact_root)
    original = indexing.index_file(source_path=str(source))
    with connect(db_path) as connection:
        connection.execute(
            "UPDATE documents SET metadata_json = ? WHERE id = ?",
            (json.dumps({"page_chunk_version": 2, "artifact_extraction_version": 7}), original.id),
        )

    def fail_chunking(_parsed):
        raise RuntimeError("synthetic chunk failure")

    monkeypatch.setattr(indexing_module, "chunk_parsed_document", fail_chunking)
    with pytest.raises(RuntimeError, match="synthetic chunk failure"):
        indexing.index_file(source_path=str(source))

    documents = indexing.list_documents()
    assert [document["id"] for document in documents] == [original.id]
    assert RagService(db_path).search("Stable evidence", top_k=1)


def test_version_reindex_preserves_document_figure_and_table_identity(tmp_path) -> None:
    import json

    from app.db.sqlite import connect

    db_path = tmp_path / "app.db"
    source = tmp_path / "paper.md"
    source.write_text(
        "# Paper\n\nFigure 3: Stable architecture diagram.\n\n"
        "| Model | Acc |\n| --- | --- |\n| Aya | 91.2 |",
        encoding="utf-8",
    )
    init_db(db_path)
    indexing = IndexingService(db_path)
    original = indexing.index_file(source_path=str(source))
    original_figure = RagService(db_path).list_document_figures(original.id)[0]
    original_table = RagService(db_path).list_document_tables(original.id)[0]
    original_chunk_ids = [
        chunk["id"] for chunk in RagService(db_path).list_document_chunks(original.id)
    ]
    with connect(db_path) as connection:
        connection.execute(
            "UPDATE documents SET metadata_json = ? WHERE id = ?",
            (json.dumps({"page_chunk_version": 2, "artifact_extraction_version": 7}), original.id),
        )

    rebuilt = indexing.index_file(source_path=str(source))
    rebuilt_figure = RagService(db_path).list_document_figures(rebuilt.id)[0]
    rebuilt_table = RagService(db_path).list_document_tables(rebuilt.id)[0]
    rebuilt_chunk_ids = [
        chunk["id"] for chunk in RagService(db_path).list_document_chunks(rebuilt.id)
    ]

    assert rebuilt.id == original.id
    assert rebuilt_figure["id"] == original_figure["id"]
    assert rebuilt_table["id"] == original_table["id"]
    assert rebuilt_chunk_ids == original_chunk_ids


def test_text_only_rechunk_preserves_existing_artifacts(tmp_path) -> None:
    db_path = tmp_path / "app.db"
    source = tmp_path / "paper.md"
    source.write_text(
        "# Paper\n\n"
        + "Method and retrieval context. " * 100
        + "\n\nFigure 3: Stable architecture diagram.\n\n"
        "| Model | Acc |\n| --- | --- |\n| Aya | 91.2 |",
        encoding="utf-8",
    )
    init_db(db_path)
    indexing = IndexingService(db_path)
    indexed = indexing.index_file(source_path=str(source))
    rag = RagService(db_path)
    figure_id = rag.list_document_figures(indexed.id)[0]["id"]
    table_id = rag.list_document_tables(indexed.id)[0]["id"]

    result = indexing.rechunk_document(indexed.id)

    assert result["chunking_version"] == 3
    assert rag.list_document_figures(indexed.id)[0]["id"] == figure_id
    assert rag.list_document_tables(indexed.id)[0]["id"] == table_id
    assert all(
        chunk["parent_chunk_id"]
        for chunk in rag.list_document_chunks(indexed.id)
    )


def test_rag_service_expands_neighbor_chunks(tmp_path) -> None:
    db_path = tmp_path / "app.db"
    docs_dir = tmp_path / "docs"
    docs_dir.mkdir()
    (docs_dir / "long.md").write_text(
        "\n\n".join(
            [
                "# Long Report",
                "Alpha architecture context. " * 80,
                "Beta retrieval target details. " * 80,
                "Gamma neighbor explanation. " * 80,
            ]
        ),
        encoding="utf-8",
    )

    init_db(db_path)
    indexing = IndexingService(db_path)
    rag = RagService(db_path)
    result = indexing.index_folder(
        folder_path=str(docs_dir),
        recursive=False,
        file_types=["md"],
    )
    chunks = rag.list_document_chunks(result.documents[0].id)
    middle = chunks[1]
    expanded = rag.expand_with_neighbor_chunks(
        [
            {
                "chunk_id": middle["id"],
                "document_id": middle["document_id"],
                "filename": middle["filename"],
                "source_path": middle["source_path"],
                "content": middle["content"],
                "chunk_index": middle["chunk_index"],
                "page_number": middle["page_number"],
                "chunk_type": "text",
            }
        ]
    )

    assert expanded[0]["expanded_content"]
    assert "retrieved chunk" in expanded[0]["expanded_content"]
    assert expanded[0]["neighbor_chunk_ids"]


def test_rag_service_expands_token_child_to_parent_context(tmp_path) -> None:
    db_path = tmp_path / "app.db"
    source = tmp_path / "parent.md"
    source.write_text(
        "# Method\n\n"
        + "Acoustic encoder setup and preprocessing details. " * 80
        + "\n\n"
        + "Cross-modal fusion target evidence and implementation details. " * 80,
        encoding="utf-8",
    )
    init_db(db_path)
    indexed = IndexingService(db_path).index_file(source_path=str(source))
    rag = RagService(db_path)
    chunks = rag.list_document_chunks(indexed.id)
    child = next(chunk for chunk in chunks if chunk["parent_chunk_id"])

    expanded = rag.expand_with_neighbor_chunks(
        [
            {
                "chunk_id": child["id"],
                "document_id": child["document_id"],
                "filename": child["filename"],
                "source_path": child["source_path"],
                "content": child["content"],
                "chunk_index": child["chunk_index"],
                "page_number": child["page_number"],
                "chunk_type": "text",
            }
        ]
    )

    assert expanded[0]["parent_expanded"] is True
    assert expanded[0]["parent_chunk_id"] == child["parent_chunk_id"]
    assert "[retrieved chunk / child]" in expanded[0]["expanded_content"]
    assert "[parent section context]" in expanded[0]["expanded_content"]
    assert isinstance(expanded[0]["neighbor_chunk_ids"], list)


def test_rag_service_neighbor_expansion_respects_section_metadata(tmp_path) -> None:
    db_path = tmp_path / "app.db"
    docs_dir = tmp_path / "docs"
    docs_dir.mkdir()
    (docs_dir / "sections.md").write_text(
        "\n\n".join(
            [
                "# Related Work",
                "Prior work context. " * 80,
                "# Method",
                "Method target details. " * 80,
                "Method continuation details. " * 80,
            ]
        ),
        encoding="utf-8",
    )

    init_db(db_path)
    indexing = IndexingService(db_path)
    rag = RagService(db_path)
    result = indexing.index_folder(
        folder_path=str(docs_dir),
        recursive=False,
        file_types=["md"],
    )
    chunks = rag.list_document_chunks(result.documents[0].id)
    method_chunk = next(chunk for chunk in chunks if chunk["metadata"].get("section_title") == "Method")
    expanded = rag.expand_with_neighbor_chunks(
        [
            {
                "chunk_id": method_chunk["id"],
                "document_id": method_chunk["document_id"],
                "filename": method_chunk["filename"],
                "source_path": method_chunk["source_path"],
                "content": method_chunk["content"],
                "chunk_index": method_chunk["chunk_index"],
                "page_number": method_chunk["page_number"],
                "chunk_type": "text",
                "section_title": "Method",
            }
        ],
        window=2,
    )

    assert "Prior work context" not in expanded[0].get("expanded_content", "")


def test_parse_docx(tmp_path) -> None:
    from docx import Document

    path = tmp_path / "note.docx"
    document = Document()
    document.add_paragraph("Docx retrieval works.")
    document.save(path)

    assert "Docx retrieval works." in parse_text_file(path)


def test_pdf_image_extraction_saves_figure_artifact(tmp_path) -> None:
    import pytest

    fitz = pytest.importorskip("fitz")
    path = tmp_path / "visual.pdf"
    document = fitz.open()
    page = document.new_page()
    pixmap = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 40, 40), 0)
    pixmap.clear_with(0x22AA66)
    page.insert_image(fitz.Rect(50, 50, 120, 120), pixmap=pixmap)
    page.insert_text((50, 140), "Figure 1: Green architecture block")
    document.save(path)

    parsed = parse_document(
        path,
        artifact_dir=tmp_path / "artifacts",
        vision_summarizer=lambda image_path, caption, page_text: f"Visual summary for {caption}",
    )
    figures = [figure for figure in parsed.figures or [] if figure.image_path]

    assert figures
    assert parsed.parser_name in {"docling", "pypdf+pymupdf"}
    assert Path(figures[0].image_path or "").exists()
    assert figures[0].visual_summary == "Visual summary for Figure 1: Green architecture block"
    assert figures[0].extraction_method in {"docling_picture", "pymupdf_caption_crop"}
    assert figures[0].bbox is not None
    assert figures[0].metadata["source"] in {"docling", "pymupdf_caption_crop"}
    assert figures[0].metadata["width"] > 40


def test_pdf_ocr_gate_is_document_level_and_not_filename_specific() -> None:
    from app.rag.parsers import ParsedPage, _needs_pdf_ocr

    scanned = [ParsedPage(page_number=index, text="") for index in range(1, 11)]
    normal = [
        ParsedPage(page_number=index, text="searchable scientific paragraph " * 8)
        for index in range(1, 8)
    ] + [ParsedPage(page_number=8, text="")]
    sparse_but_usable = [
        ParsedPage(page_number=index, text="searchable content " * 4)
        for index in range(1, 5)
    ] + [ParsedPage(page_number=index, text="") for index in range(5, 11)]

    assert _needs_pdf_ocr(scanned) is True
    assert _needs_pdf_ocr(normal) is False
    # Four usable pages out of ten are enough to preserve the fast text path.
    assert _needs_pdf_ocr(sparse_but_usable) is False


def test_page_ocr_fallback_keeps_searchable_pages_and_page_provenance(
    tmp_path, monkeypatch
) -> None:
    import sys
    from types import ModuleType, SimpleNamespace

    from app.rag.parsers import ParsedPage, _ocr_sparse_pdf_pages

    class FakePage:
        def get_pixmap(self, **kwargs):  # noqa: ANN003
            return SimpleNamespace(tobytes=lambda _: b"page-image")

    class FakeDocument:
        def __enter__(self):
            return self

        def __exit__(self, *args):  # noqa: ANN002
            return None

        def __iter__(self):
            return iter([FakePage(), FakePage()])

    fake_fitz = ModuleType("fitz")
    fake_fitz.open = lambda _: FakeDocument()
    fake_fitz.Matrix = lambda *_: object()

    class FakeRapidOCR:
        def __init__(self, **kwargs):  # noqa: ANN003
            pass

        def __call__(self, payload):  # noqa: ANN001
            return SimpleNamespace(txts=("OCR title", "OCR value 42"))

    fake_rapidocr = ModuleType("rapidocr")
    fake_rapidocr.RapidOCR = FakeRapidOCR
    fake_typings = ModuleType("rapidocr.utils.typings")
    fake_typings.EngineType = SimpleNamespace(TORCH="torch")
    monkeypatch.setitem(sys.modules, "fitz", fake_fitz)
    monkeypatch.setitem(sys.modules, "rapidocr", fake_rapidocr)
    monkeypatch.setitem(sys.modules, "rapidocr.utils.typings", fake_typings)

    pages = [
        ParsedPage(1, "existing searchable text " * 4),
        ParsedPage(2, ""),
    ]
    result = _ocr_sparse_pdf_pages(tmp_path / "scanned.pdf", pages)

    assert result[0] == pages[0]
    assert result[1].page_number == 2
    assert result[1].text == "OCR title\nOCR value 42"


def test_figure_caption_requires_separator(tmp_path) -> None:
    path = tmp_path / "captions.md"
    path.write_text(
        "\n".join(
            [
                "Figure 2 complements Table 4 with additional results.",
                "Figure 3: Real architecture caption",
                "Fig. 4. Another real caption",
            ]
        ),
        encoding="utf-8",
    )

    parsed = parse_document(path)
    captions = [figure.caption for figure in parsed.figures or []]

    assert captions == ["Figure 3: Real architecture caption", "Fig. 4. Another real caption"]


def test_figure_caption_merges_continuation_lines(tmp_path) -> None:
    path = tmp_path / "multiline-caption.md"
    path.write_text(
        "\n".join(
            [
                "Fig. 3. The key-sparse attention in KS-Transformer.",
                "In this module, only top-k keys are kept.",
                "Fig. 4. The details of CCAB.",
            ]
        ),
        encoding="utf-8",
    )

    parsed = parse_document(path)
    captions = [figure.caption for figure in parsed.figures or []]

    assert captions == [
        "Fig. 3. The key-sparse attention in KS-Transformer. In this module, only top-k keys are kept.",
        "Fig. 4. The details of CCAB.",
    ]
